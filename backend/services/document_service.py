import os
import sys
from typing import List, Optional
from pathlib import Path
import logging

# Ensure correct site-packages path
venv_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.venv', 'lib', 'python3.12', 'site-packages')
if venv_path not in sys.path:
    sys.path.insert(0, venv_path)

from config import OPENAI_API_KEY, OPENAI_EMBEDDING_MODEL, CHROMA_DIR, UPLOAD_DIR

logger = logging.getLogger(__name__)

# Try to import optional dependencies
LANGCHAIN_AVAILABLE = False
Chroma = None
OpenAIEmbeddings = None
RecursiveCharacterTextSplitter = None
PyPDFLoader = None
TextLoader = None

try:
    import langchain_text_splitters
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    logger.info("langchain_text_splitters loaded")
except ImportError as e:
    logger.warning(f"langchain_text_splitters not available: {e}")

try:
    import langchain_community
    from langchain_community.document_loaders import PyPDFLoader, TextLoader
    logger.info("langchain_community loaded")
except ImportError as e:
    logger.warning(f"langchain_community not available: {e}")

try:
    import langchain_chroma
    from langchain_chroma import Chroma
    logger.info("langchain_chroma loaded")
except ImportError as e:
    logger.warning(f"langchain_chroma not available: {e}")

try:
    import langchain_openai
    from langchain_openai import OpenAIEmbeddings
    logger.info("langchain_openai loaded")
except ImportError as e:
    logger.warning(f"langchain_openai not available: {e}")

# Check if all required components are available
if all([RecursiveCharacterTextSplitter, PyPDFLoader, TextLoader, Chroma, OpenAIEmbeddings]):
    LANGCHAIN_AVAILABLE = True
    logger.info("All LangChain dependencies loaded successfully")
else:
    logger.warning("Some LangChain dependencies missing. Using fallback document processing.")

class DocumentService:
    """Service for processing and indexing documents for RAG using OpenAI embeddings."""
    
    def __init__(self, persist_directory: str = None):
        self.persist_directory = persist_directory or CHROMA_DIR
        self.upload_directory = UPLOAD_DIR
        self.vector_store = None
        self.documents = []
        self.simple_store = []  # Always initialize simple_store as fallback
        
        # Ensure directories exist
        os.makedirs(self.persist_directory, exist_ok=True)
        os.makedirs(self.upload_directory, exist_ok=True)
        
        if LANGCHAIN_AVAILABLE:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
            self._initialize_vector_store()
        else:
            self.text_splitter = None
            logger.info("Using simple in-memory document store")
    
    def _initialize_vector_store(self):
        """Initialize or load the vector store with OpenAI embeddings."""
        try:
            # Use OpenAI text-embedding-3-large model
            self.embeddings = OpenAIEmbeddings(
                model=OPENAI_EMBEDDING_MODEL,
                openai_api_key=OPENAI_API_KEY
            )
            
            # Initialize Chroma vector store
            self.vector_store = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings,
                collection_name="electrolyte_docs"
            )
            logger.info(f"Vector store initialized with {OPENAI_EMBEDDING_MODEL}")
        except Exception as e:
            logger.error(f"Failed to initialize vector store: {e}")
            import traceback
            logger.error(traceback.format_exc())
            self._initialize_simple_store()
    
    def _initialize_simple_store(self):
        """Initialize a simple in-memory document store as fallback."""
        # simple_store is already initialized in __init__
        logger.info("Using simple in-memory document store")
    
    async def process_pdf(self, file_path: str, filename: str) -> dict:
        """Process a PDF file and add to vector store."""
        logger.info(f"Processing PDF: {filename} from {file_path}")
        
        # Verify file exists
        if not os.path.exists(file_path):
            logger.error(f"PDF file not found: {file_path}")
            return {
                "filename": filename,
                "status": "error",
                "chunks_created": 0,
                "message": f"File not found: {file_path}"
            }
        
        try:
            if LANGCHAIN_AVAILABLE and self.vector_store and PyPDFLoader:
                logger.info(f"Using LangChain PDF loader for {filename}")
                return await self._process_pdf_langchain(file_path, filename)
            else:
                logger.info(f"Using simple PDF loader for {filename}")
                return await self._process_pdf_simple(file_path, filename)
        except Exception as e:
            logger.error(f"Error processing PDF with primary method: {e}")
            # Try fallback to simple processing
            try:
                logger.info(f"Falling back to simple PDF processing for {filename}")
                return await self._process_pdf_simple(file_path, filename)
            except Exception as fallback_error:
                logger.error(f"Fallback PDF processing also failed: {fallback_error}")
                return {
                    "filename": filename,
                    "status": "error",
                    "chunks_created": 0,
                    "message": f"PDF processing failed: {str(e)}. Fallback also failed: {str(fallback_error)}"
                }
    
    async def _process_pdf_langchain(self, file_path: str, filename: str) -> dict:
        """Process PDF using LangChain with OpenAI embeddings."""
        try:
            logger.info(f"Loading PDF with PyPDFLoader: {file_path}")
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            logger.info(f"Loaded {len(pages)} pages from {filename}")
            
            if not pages:
                raise Exception("No pages extracted from PDF")
            
            # Split into chunks
            chunks = self.text_splitter.split_documents(pages)
            logger.info(f"Split into {len(chunks)} chunks")
            
            if not chunks:
                raise Exception("No chunks created from PDF")
            
            # Add metadata
            for chunk in chunks:
                chunk.metadata["source"] = filename
            
            # Add to vector store (embeddings via text-embedding-3-large)
            self.vector_store.add_documents(chunks)
            logger.info(f"Added {len(chunks)} chunks to vector store")
            
            return {
                "filename": filename,
                "status": "success",
                "chunks_created": len(chunks),
                "message": f"Successfully indexed {len(chunks)} chunks from {filename} using {OPENAI_EMBEDDING_MODEL}"
            }
        except Exception as e:
            logger.error(f"LangChain PDF processing failed for {filename}: {e}")
            raise Exception(f"LangChain processing failed: {e}")
    
    async def _process_pdf_simple(self, file_path: str, filename: str) -> dict:
        """Process PDF using simple extraction with pypdf."""
        try:
            from pypdf import PdfReader
            
            logger.info(f"Reading PDF with pypdf: {file_path}")
            reader = PdfReader(file_path)
            text_chunks = []
            total_text = ""
            
            logger.info(f"PDF has {len(reader.pages)} pages")
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        total_text += text + "\n"
                        # Simple chunking by words
                        words = text.split()
                        chunk_size = 200  # words per chunk
                        for i in range(0, len(words), chunk_size):
                            chunk_text = " ".join(words[i:i + chunk_size])
                            if chunk_text.strip():
                                text_chunks.append({
                                    "content": chunk_text,
                                    "source": filename,
                                    "page": page_num + 1
                                })
                except Exception as page_error:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {page_error}")
                    continue
            
            if not text_chunks:
                logger.warning(f"No text extracted from PDF: {filename}")
                return {
                    "filename": filename,
                    "status": "warning",
                    "chunks_created": 0,
                    "message": f"No text content could be extracted from {filename}. The PDF may be scanned/image-based."
                }
            
            # Store in simple store and also try to add to vector store if available
            self.simple_store.extend(text_chunks)
            self.documents.extend(text_chunks)
            
            # If vector store is available, also add to it
            if LANGCHAIN_AVAILABLE and self.vector_store:
                try:
                    from langchain_core.documents import Document as LCDocument
                    lc_docs = [
                        LCDocument(page_content=chunk["content"], metadata={"source": filename, "page": chunk.get("page", 1)})
                        for chunk in text_chunks
                    ]
                    self.vector_store.add_documents(lc_docs)
                    logger.info(f"Added {len(lc_docs)} chunks to vector store via simple processing")
                except Exception as vs_error:
                    logger.warning(f"Could not add to vector store: {vs_error}")
            
            logger.info(f"Successfully processed {len(text_chunks)} chunks from {filename}")
            
            return {
                "filename": filename,
                "status": "success",
                "chunks_created": len(text_chunks),
                "message": f"Successfully processed {len(text_chunks)} chunks from {filename}"
            }
        except Exception as e:
            logger.error(f"Simple PDF processing failed for {filename}: {e}")
            raise Exception(f"Simple PDF processing failed: {e}")
    
    async def process_text(self, file_path: str, filename: str) -> dict:
        """Process a text file and add to vector store."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if LANGCHAIN_AVAILABLE and self.vector_store:
                loader = TextLoader(file_path)
                documents = loader.load()
                chunks = self.text_splitter.split_documents(documents)
                
                for chunk in chunks:
                    chunk.metadata["source"] = filename
                
                self.vector_store.add_documents(chunks)
                
                return {
                    "filename": filename,
                    "status": "success",
                    "chunks_created": len(chunks),
                    "message": f"Successfully indexed {len(chunks)} chunks from {filename} using {OPENAI_EMBEDDING_MODEL}"
                }
            else:
                # Simple processing
                words = content.split()
                chunk_size = 200
                chunks = []
                
                for i in range(0, len(words), chunk_size):
                    chunk_text = " ".join(words[i:i + chunk_size])
                    chunks.append({
                        "content": chunk_text,
                        "source": filename
                    })
                
                self.simple_store.extend(chunks)
                self.documents.extend(chunks)
                
                return {
                    "filename": filename,
                    "status": "success",
                    "chunks_created": len(chunks),
                    "message": f"Successfully processed {len(chunks)} chunks from {filename}"
                }
        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            return {
                "filename": filename,
                "status": "error",
                "chunks_created": 0,
                "message": str(e)
            }
    
    async def process_docx(self, file_path: str, filename: str) -> dict:
        """Process a DOCX file and add to vector store."""
        try:
            from docx import Document
            
            # Extract text from DOCX
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            content = "\n".join(full_text)
            
            if not content.strip():
                return {
                    "filename": filename,
                    "status": "error",
                    "chunks_created": 0,
                    "message": "No text content found in DOCX file"
                }
            
            if LANGCHAIN_AVAILABLE and self.vector_store and self.text_splitter:
                # Create document objects for LangChain
                from langchain_core.documents import Document as LCDocument
                documents = [LCDocument(page_content=content, metadata={"source": filename})]
                chunks = self.text_splitter.split_documents(documents)
                
                for chunk in chunks:
                    chunk.metadata["source"] = filename
                
                self.vector_store.add_documents(chunks)
                
                return {
                    "filename": filename,
                    "status": "success",
                    "chunks_created": len(chunks),
                    "message": f"Successfully indexed {len(chunks)} chunks from {filename} using {OPENAI_EMBEDDING_MODEL}"
                }
            else:
                # Simple processing
                words = content.split()
                chunk_size = 200
                chunks = []
                
                for i in range(0, len(words), chunk_size):
                    chunk_text = " ".join(words[i:i + chunk_size])
                    chunks.append({
                        "content": chunk_text,
                        "source": filename
                    })
                
                self.simple_store.extend(chunks)
                self.documents.extend(chunks)
                
                return {
                    "filename": filename,
                    "status": "success",
                    "chunks_created": len(chunks),
                    "message": f"Successfully processed {len(chunks)} chunks from {filename}"
                }
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            return {
                "filename": filename,
                "status": "error",
                "chunks_created": 0,
                "message": str(e)
            }
    
    def search(self, query: str, k: int = 5) -> List[dict]:
        """Search the document store using OpenAI embeddings."""
        if LANGCHAIN_AVAILABLE and self.vector_store:
            try:
                results = self.vector_store.similarity_search(query, k=k)
                return [
                    {
                        "content": doc.page_content,
                        "source": doc.metadata.get("source", "unknown"),
                        "score": 0.85
                    }
                    for doc in results
                ]
            except Exception as e:
                logger.error(f"Vector search failed: {e}")
                return self._simple_search(query, k)
        else:
            return self._simple_search(query, k)
    
    def _simple_search(self, query: str, k: int = 5) -> List[dict]:
        """Simple keyword-based search."""
        query_words = set(query.lower().split())
        scored_docs = []
        
        for doc in self.simple_store:
            content = doc.get("content", "").lower()
            doc_words = set(content.split())
            
            # Calculate simple overlap score
            overlap = len(query_words & doc_words)
            if overlap > 0:
                scored_docs.append({
                    "content": doc.get("content", ""),
                    "source": doc.get("source", "unknown"),
                    "score": overlap / len(query_words)
                })
        
        # Sort by score and return top k
        scored_docs.sort(key=lambda x: x["score"], reverse=True)
        return scored_docs[:k]
    
    def get_vector_store(self):
        """Get the vector store for use by agents."""
        return self.vector_store
    
    def get_document_count(self) -> int:
        """Get total number of indexed documents."""
        if LANGCHAIN_AVAILABLE and self.vector_store:
            try:
                return len(self.vector_store.get()['ids'])
            except:
                return 0
        return len(self.simple_store)
    
    def clear_store(self):
        """Clear all documents from the store."""
        if LANGCHAIN_AVAILABLE and self.vector_store:
            try:
                # Delete and recreate
                self.vector_store.delete_collection()
                self._initialize_vector_store()
            except Exception as e:
                logger.error(f"Failed to clear vector store: {e}")
        
        # Always clear in-memory stores
        self.simple_store = []
        self.documents = []
        logger.info("Document store cleared")
